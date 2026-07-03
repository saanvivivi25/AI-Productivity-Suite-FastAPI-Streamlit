from __future__ import annotations

# =====================================================================
# IMPORTS
# =====================================================================
import datetime
import logging
import sqlite3
import sys
import threading
from functools import lru_cache
from docx import Document
from io import BytesIO
from typing import Optional

import google.generativeai as genai
import pandas as pd
from fastapi import Depends, FastAPI, File, Query, Request, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
from pydantic_settings import BaseSettings, SettingsConfigDict


# =====================================================================
# CONFIGURATION  (equivalent to app/core/config.py)
# =====================================================================
class Settings(BaseSettings):
    # --- App ---
    APP_NAME: str = "AI Productivity Suite API"
    APP_VERSION: str = "1.0.0"
    DEBUG: bool = False

    # --- Gemini ---
    GEMINI_API_KEY: str = "API"
    GEMINI_MODEL: str = "gemini-2.5-flash"

    # --- Database ---
    DATABASE_PATH: str = "ai_productivity.db"

    # --- Logging ---
    LOG_LEVEL: str = "INFO"
    LOG_FILE: str = "app.log"

    model_config = SettingsConfigDict(
        env_file=".env",
        env_file_encoding="utf-8",
        case_sensitive=True,
        extra="ignore",
    )


@lru_cache
def get_settings() -> Settings:
    """Cached settings instance so the .env file is parsed only once."""
    return Settings()


# =====================================================================
# LOGGING  (equivalent to app/core/logging_config.py)
# =====================================================================
def setup_logging() -> None:
    settings = get_settings()
    log_level = getattr(logging, settings.LOG_LEVEL.upper(), logging.INFO)

    formatter = logging.Formatter(
        fmt="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    root_logger = logging.getLogger()
    root_logger.setLevel(log_level)

    if root_logger.handlers:
        return

    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setFormatter(formatter)
    root_logger.addHandler(console_handler)

    file_handler = logging.FileHandler(settings.LOG_FILE, encoding="utf-8")
    file_handler.setFormatter(formatter)
    root_logger.addHandler(file_handler)

    logging.getLogger("uvicorn.access").setLevel(logging.WARNING)


setup_logging()
logger = logging.getLogger(__name__)


# =====================================================================
# EXCEPTIONS  (equivalent to app/core/exceptions.py)
# =====================================================================
class AppException(Exception):
    """Base class for all application-specific errors."""

    status_code = 500
    message = "Internal server error."

    def __init__(self, message: Optional[str] = None):
        self.message = message or self.message
        super().__init__(self.message)


class GeminiServiceError(AppException):
    status_code = 502
    message = "The AI service failed to generate a response."


class EmptyInputError(AppException):
    status_code = 422
    message = "Input text must not be empty."


class RecordNotFoundError(AppException):
    status_code = 404
    message = "Record not found."


class InvalidFileError(AppException):
    status_code = 400
    message = "Uploaded file is invalid."


def register_exception_handlers(app: FastAPI) -> None:
    @app.exception_handler(AppException)
    async def app_exception_handler(request: Request, exc: AppException):
        logger.error("AppException on %s %s: %s", request.method, request.url.path, exc.message)
        return JSONResponse(
            status_code=exc.status_code,
            content={"error": exc.__class__.__name__, "detail": exc.message},
        )

    @app.exception_handler(Exception)
    async def unhandled_exception_handler(request: Request, exc: Exception):
        logger.exception("Unhandled exception on %s %s", request.method, request.url.path)
        return JSONResponse(
            status_code=500,
            content={"error": "InternalServerError", "detail": "Something went wrong."},
        )


# =====================================================================
# DATABASE  (equivalent to app/db/database.py)
# =====================================================================
class HistoryRepository:
    """Data access object for the `history` table."""

    def __init__(self, settings: Settings):
        self._db_path = settings.DATABASE_PATH
        self._lock = threading.Lock()
        self._conn = sqlite3.connect(self._db_path, check_same_thread=False)
        self._create_table()

    def _create_table(self) -> None:
        with self._lock:
            self._conn.execute(
                """
                CREATE TABLE IF NOT EXISTS history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    module TEXT,
                    input_text TEXT,
                    output_text TEXT,
                    extra TEXT,
                    date TEXT
                )
                """
            )
            self._conn.commit()
        logger.info("Ensured 'history' table exists at %s", self._db_path)

    def save(self, module: str, input_text: str, output_text: str, extra: str = "") -> int:
        with self._lock:
            cursor = self._conn.execute(
                """
                INSERT INTO history (module, input_text, output_text, extra, date)
                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    module,
                    input_text,
                    output_text,
                    extra,
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                ),
            )
            self._conn.commit()
            record_id = cursor.lastrowid
        logger.debug("Saved history record id=%s module=%s", record_id, module)
        return record_id

    def load_all(self) -> pd.DataFrame:
        return pd.read_sql("SELECT * FROM history ORDER BY id DESC", self._conn)

    def search(self, keyword: str) -> pd.DataFrame:
        query = """
            SELECT * FROM history
            WHERE module LIKE ?
               OR input_text LIKE ?
               OR output_text LIKE ?
               OR extra LIKE ?
            ORDER BY id DESC
        """
        key = f"%{keyword}%"
        return pd.read_sql(query, self._conn, params=(key, key, key, key))

    def delete(self, record_id: int) -> bool:
        with self._lock:
            cursor = self._conn.execute("DELETE FROM history WHERE id=?", (record_id,))
            self._conn.commit()
            deleted = cursor.rowcount > 0
        logger.debug("Delete record id=%s -> %s", record_id, deleted)
        return deleted

    def delete_all(self) -> None:
        with self._lock:
            self._conn.execute("DELETE FROM history")
            self._conn.commit()
        logger.info("Deleted all history records")

    @staticmethod
    def export_word(df: pd.DataFrame) -> BytesIO:

        document = Document()

        document.add_heading("AI Productivity History Report", level=1)

        if df.empty:
            document.add_paragraph("No history records found.")

        else:
            for _, row in df.iterrows():

                document.add_heading(
                    f"Record ID: {row['id']}",
                    level=2
                )

                document.add_paragraph(f"Module: {row['module']}")

                document.add_heading("Input", level=3)
                document.add_paragraph(str(row["input_text"]))

                document.add_heading("Output", level=3)
                document.add_paragraph(str(row["output_text"]))

                if str(row["extra"]).strip():
                    document.add_heading("Extra", level=3)
                    document.add_paragraph(str(row["extra"]))

                document.add_paragraph(
                    f"Date: {row['date']}"
                )

                document.add_page_break()

        output = BytesIO()

        document.save(output)

        output.seek(0)

        return output


# =====================================================================
# SCHEMAS  (equivalent to app/models/schemas.py)
# =====================================================================
class TextRequest(BaseModel):
    text: str = Field(..., min_length=1)


class ResumeRequest(BaseModel):
    resume: str = Field(..., min_length=1)


class MeetingTranscriptRequest(BaseModel):
    transcript: str = Field(..., min_length=1)


class FeedbackRequest(BaseModel):
    feedback: str = Field(..., min_length=1)


class EmailRequest(BaseModel):
    subject: str = Field(..., min_length=1)
    purpose: str = Field(..., min_length=1)


class SummarizeRequest(BaseModel):
    text: str = Field(..., min_length=1)


class AIResultResponse(BaseModel):
    result: str
    history_id: int


class HistoryRecord(BaseModel):
    id: int
    module: str
    input_text: str
    output_text: str
    extra: Optional[str] = ""
    date: str


class HistoryListResponse(BaseModel):
    count: int
    records: list[HistoryRecord]


class DeleteResponse(BaseModel):
    detail: str


class DashboardStatsResponse(BaseModel):
    total_records: int
    stats_by_module: dict[str, int]


# =====================================================================
# SERVICES  (equivalent to app/services/*.py)
# =====================================================================
class GeminiService:
    """Thin wrapper around google.generativeai."""

    def __init__(self, settings: Settings):
        genai.configure(api_key=settings.GEMINI_API_KEY)
        self._model = genai.GenerativeModel(settings.GEMINI_MODEL)

    def ask(self, prompt: str) -> str:
        try:
            response = self._model.generate_content(prompt)
        except Exception as exc:  # noqa: BLE001
            logger.exception("Gemini API call failed")
            raise GeminiServiceError(f"Gemini API call failed: {exc}") from exc

        if getattr(response, "text", None):
            return response.text

        logger.warning("Gemini returned an empty response")
        raise GeminiServiceError("No response generated by the AI model.")


class AIService:
    """Business logic / prompt-building for each productivity module."""

    def __init__(self, gemini_service: GeminiService):
        self._gemini = gemini_service

    def analyze_resume(self, resume: str) -> str:
        prompt = f"""
You are an HR Recruiter.

Analyze the following resume.

Return:

Candidate Name

Skills

Experience

Strengths

Weaknesses

Recommendation
(Strong Hire / Consider / Reject)

Resume:

{resume}
"""
        logger.info("Analyzing resume (%d chars)", len(resume))
        return self._gemini.ask(prompt)

    def generate_meeting_notes(self, transcript: str) -> str:
        prompt = f"""
Convert the following meeting transcript into professional notes.

Include:

Summary

Key Discussion

Action Items

Decisions

Transcript:

{transcript}
"""
        logger.info("Generating meeting notes (%d chars)", len(transcript))
        return self._gemini.ask(prompt)

    def analyze_feedback(self, feedback: str) -> str:
        prompt = f"""
Analyze this customer feedback.

Return:

Sentiment

Category

Summary

Feedback:

{feedback}
"""
        logger.info("Analyzing feedback (%d chars)", len(feedback))
        return self._gemini.ask(prompt)

    def generate_email(self, subject: str, purpose: str) -> str:
        prompt = f"""
Write a professional email.

Subject:

{subject}

Purpose:

{purpose}
"""
        logger.info("Generating email for subject=%r", subject)
        return self._gemini.ask(prompt)

    def summarize_text(self, text: str) -> str:
        prompt = f"""
Summarize the following text.

Use bullet points.

Text:

{text}
"""
        logger.info("Summarizing text (%d chars)", len(text))
        return self._gemini.ask(prompt)


class HistoryService:
    """Business logic layer over HistoryRepository."""

    def __init__(self, repository: HistoryRepository):
        self._repo = repository

    def save(self, module: str, input_text: str, output_text: str, extra: str = "") -> int:
        return self._repo.save(module, input_text, output_text, extra)

    def list_all(self) -> list[HistoryRecord]:
        return self._df_to_records(self._repo.load_all())

    def search(self, keyword: str) -> list[HistoryRecord]:
        return self._df_to_records(self._repo.search(keyword))

    def delete(self, record_id: int) -> None:
        if not self._repo.delete(record_id):
            raise RecordNotFoundError(f"No history record with id={record_id}")
        logger.info("Deleted history record id=%s", record_id)

    def delete_all(self) -> None:
        self._repo.delete_all()

    def export_word(self, keyword: Optional[str] = None) -> BytesIO:
        df = self._repo.search(keyword) if keyword else self._repo.load_all()
        return self._repo.export_word(df)

    def dashboard_stats(self) -> DashboardStatsResponse:
        df = self._repo.load_all()
        total = len(df)
        stats = {} if df.empty else df["module"].value_counts().to_dict()
        return DashboardStatsResponse(total_records=total, stats_by_module=stats)

    @staticmethod
    def _df_to_records(df: pd.DataFrame) -> list[HistoryRecord]:
        if df.empty:
            return []
        df = df.fillna("")
        return [HistoryRecord(**row) for row in df.to_dict(orient="records")]


# =====================================================================
# DEPENDENCY INJECTION  (equivalent to app/dependencies.py)
# =====================================================================
@lru_cache
def get_history_repository() -> HistoryRepository:
    return HistoryRepository(get_settings())


@lru_cache
def get_gemini_service() -> GeminiService:
    return GeminiService(get_settings())


@lru_cache
def get_ai_service() -> AIService:
    return AIService(get_gemini_service())


@lru_cache
def get_history_service() -> HistoryService:
    return HistoryService(get_history_repository())


# =====================================================================
# APP SETUP
# =====================================================================
settings = get_settings()

app = FastAPI(
    title=settings.APP_NAME,
    version=settings.APP_VERSION,
    debug=settings.DEBUG,
)

register_exception_handlers(app)


@app.on_event("startup")
async def on_startup():
    logger.info("%s v%s starting up", settings.APP_NAME, settings.APP_VERSION)


@app.get("/", tags=["Health"])
def root():
    return {"status": "ok", "app": settings.APP_NAME, "version": settings.APP_VERSION}


@app.get("/health", tags=["Health"])
def health():
    return {"status": "healthy"}


# =====================================================================
# ROUTES — DASHBOARD  (equivalent to app/routers/dashboard.py)
# =====================================================================
@app.get("/dashboard/stats", response_model=DashboardStatsResponse, tags=["Dashboard"])
def get_dashboard_stats(history_service: HistoryService = Depends(get_history_service)):
    return history_service.dashboard_stats()


# =====================================================================
# ROUTES — RESUME SCREENING  (equivalent to app/routers/resume.py)
# =====================================================================
RESUME_MODULE = "Resume Screening"


@app.post("/resume/analyze-bulk", tags=["Resume Screening"])
async def analyze_resumes_bulk(
    file: UploadFile = File(..., description="Word file with resumes in the first column"),
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not file.filename.endswith((".docx" , ".doc")):
        raise InvalidFileError("Please upload a Word (.docx) file.")

    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents), engine="openpyxl")
    except Exception as exc:
        raise InvalidFileError(f"Could not read word file: {exc}") from exc

    outputs = []

    for resume in df.iloc[:, 0]:
        result = ai_service.analyze_resume(str(resume))
        outputs.append(result)

        history_service.save(
            RESUME_MODULE,
            str(resume),
            result
        )

    # -------------------------------
    # Create Word Document
    # -------------------------------

    document = Document()

    document.add_heading("Resume Analysis Report", level=1)

    for i, resume in enumerate(df.iloc[:, 0]):

        document.add_heading(f"Resume {i + 1}", level=2)

        document.add_heading("Resume", level=3)
        document.add_paragraph(str(resume))

        document.add_heading("AI Analysis", level=3)
        document.add_paragraph(outputs[i])

        if i != len(df) - 1:
            document.add_page_break()

    doc_bytes = BytesIO()

    document.save(doc_bytes)

    doc_bytes.seek(0)

    return StreamingResponse(
        doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=Resume_Analysis_Report.docx"
        },
    )

# =====================================================================
# ROUTES — MEETING NOTES  (equivalent to app/routers/meeting.py)
# =====================================================================
MEETING_MODULE = "Meeting Notes"


@app.post("/meeting-notes/generate", response_model=AIResultResponse, tags=["Meeting Notes"])
def generate_meeting_notes(
    payload: MeetingTranscriptRequest,
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not payload.transcript.strip():
        raise EmptyInputError("Transcript must not be empty.")

    result = ai_service.generate_meeting_notes(payload.transcript)
    history_id = history_service.save(MEETING_MODULE, payload.transcript, result)
    return AIResultResponse(result=result, history_id=history_id)


# =====================================================================
# ROUTES — FEEDBACK ANALYZER  (equivalent to app/routers/feedback.py)
# =====================================================================
FEEDBACK_MODULE = "Feedback Analyzer"


@app.post("/feedback/analyze", response_model=AIResultResponse, tags=["Feedback Analyzer"])
def analyze_feedback(
    payload: FeedbackRequest,
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not payload.feedback.strip():
        raise EmptyInputError("Feedback text must not be empty.")

    result = ai_service.analyze_feedback(payload.feedback)
    history_id = history_service.save(FEEDBACK_MODULE, payload.feedback, result)
    return AIResultResponse(result=result, history_id=history_id)


@app.post("/feedback/analyze-bulk", tags=["Feedback Analyzer"])
async def analyze_feedback_bulk(
    file: UploadFile = File(..., description="Excel file with feedback in the first column"),
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not file.filename.endswith((".xlsx", ".xls")):
        raise InvalidFileError("Please upload an Excel (.xlsx/.xls) file.")

    try:
        contents = await file.read()
        df = pd.read_excel(contents)
    except Exception as exc:  # noqa: BLE001
        raise InvalidFileError(f"Could not read Excel file: {exc}") from exc

    outputs = []
    for feedback in df.iloc[:, 0]:
        result = ai_service.analyze_feedback(str(feedback))
        outputs.append(result)
        history_service.save(FEEDBACK_MODULE, str(feedback), result)

    df["AI Analysis"] = outputs
    excel_bytes = HistoryRepository.export_excel(df)
    return StreamingResponse(
        excel_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=feedback_analysis.xlsx"},
    )


# =====================================================================
# ROUTES — EMAIL GENERATOR  (equivalent to app/routers/email.py)
# =====================================================================
EMAIL_MODULE = "Email Generator"


@app.post("/email/generate", response_model=AIResultResponse, tags=["Email Generator"])
def generate_email(
    payload: EmailRequest,
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not payload.subject.strip() or not payload.purpose.strip():
        raise EmptyInputError("Both subject and purpose must be provided.")

    result = ai_service.generate_email(payload.subject, payload.purpose)
    history_id = history_service.save(EMAIL_MODULE, payload.purpose, result, payload.subject)
    return AIResultResponse(result=result, history_id=history_id)


# =====================================================================
# ROUTES — TEXT SUMMARIZER  (equivalent to app/routers/summarizer.py)
# =====================================================================
SUMMARIZER_MODULE = "Text Summarizer"


@app.post("/summarizer/summarize", response_model=AIResultResponse, tags=["Text Summarizer"])
def summarize_text(
    payload: SummarizeRequest,
    ai_service: AIService = Depends(get_ai_service),
    history_service: HistoryService = Depends(get_history_service),
):
    if not payload.text.strip():
        raise EmptyInputError("Text must not be empty.")

    result = ai_service.summarize_text(payload.text)
    history_id = history_service.save(SUMMARIZER_MODULE, payload.text, result)
    return AIResultResponse(result=result, history_id=history_id)


# =====================================================================
# ROUTES — HISTORY  (equivalent to app/routers/history.py)
# =====================================================================
@app.get("/history", response_model=HistoryListResponse, tags=["History"])
def get_history(
    search: Optional[str] = Query(None, description="Optional keyword to filter records"),
    history_service: HistoryService = Depends(get_history_service),
):
    records = history_service.search(search) if search else history_service.list_all()
    return HistoryListResponse(count=len(records), records=records)


@app.get("/history/export", tags=["History"])
def export_history(
    search: Optional[str] = Query(None),
    history_service: HistoryService = Depends(get_history_service),
):
    word_bytes = history_service.export_word(search)

    return StreamingResponse(
        word_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=AI_Productivity_Report.docx"
        },
    )
    # Get history records
    records = history_service.get_history(search)

    document = Document()

    document.add_heading("AI Productivity History Report", level=1)

    if not records:
        document.add_paragraph("No records found.")

    else:

        for record in records:

            document.add_heading(
                f"Record ID: {record['id']}",
                level=2,
            )

            document.add_paragraph(f"Module: {record['module']}")

            document.add_paragraph(f"Created At: {record['created_at']}")

            document.add_heading("Input", level=3)
            document.add_paragraph(record["input"])

            document.add_heading("Output", level=3)
            document.add_paragraph(record["output"])

            document.add_page_break()

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=AI_Productivity_Report.docx"
        },
    )


@app.delete("/history/{record_id}", response_model=DeleteResponse, tags=["History"])
def delete_record(record_id: int, history_service: HistoryService = Depends(get_history_service)):
    history_service.delete(record_id)
    return DeleteResponse(detail=f"Record {record_id} deleted.")


@app.delete("/history", response_model=DeleteResponse, tags=["History"])
def delete_all_records(history_service: HistoryService = Depends(get_history_service)):
    history_service.delete_all()
    return DeleteResponse(detail="All records deleted.")